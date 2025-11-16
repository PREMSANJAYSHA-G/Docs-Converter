# app.py
from flask import Flask, render_template, request, send_file
import os
import tempfile
import zipfile
import shutil
from werkzeug.utils import secure_filename
from docx import Document
import pdfplumber
from fpdf import FPDF
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

# -------------------------------
# Initialize Flask application
# -------------------------------
app = Flask(__name__)

# -------------------------------
# Function: Convert DOCX to PDF
# -------------------------------
def docx_to_pdf(input_docx_path, output_pdf_path):
    """
    Converts a Word file to PDF using reportlab (works on all OS)
    """
    doc = Document(input_docx_path)
    c = canvas.Canvas(output_pdf_path, pagesize=LETTER)
    width, height = LETTER
    y = height - 40  # top margin

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Split text into lines to avoid overflow
            lines = text.split('\n')
            for line in lines:
                c.drawString(40, y, line)
                y -= 15
                if y < 40:  # bottom margin
                    c.showPage()
                    y = height - 40

    c.save()
    return output_pdf_path

# -------------------------------
# Route: Home page
# -------------------------------
@app.route("/")
def index():
    """
    Render the main page with file upload UI.
    """
    return render_template("index.html")

# -------------------------------
# Route: Convert files
# -------------------------------
@app.route("/convert", methods=["POST"])
def convert_files():
    """
    Handles uploaded files and performs conversion:
      - Word (.docx) → PDF
      - PDF → Word (.docx)
    Returns a single file or ZIP if multiple files.
    """
    if "files[]" not in request.files:
        return "No files uploaded", 400

    files = request.files.getlist("files[]")
    converted_files = []

    # Temporary directories for Word → PDF conversion
    temp_dir = tempfile.mkdtemp()

    try:
        for file in files:
            if file.filename == "":
                continue

            # Sanitize filename
            filename = secure_filename(file.filename)
            ext = filename.rsplit(".", 1)[1].lower()
            temp_input = os.path.join(temp_dir, filename)
            file.save(temp_input)  # Save uploaded file temporarily

            # ---------------------------
            # Word → PDF
            # ---------------------------
            if ext == "docx":
                temp_pdf = os.path.join(temp_dir, filename.replace(".docx", ".pdf"))
                docx_to_pdf(temp_input, temp_pdf)
                converted_files.append(temp_pdf)

            # ---------------------------
            # PDF → Word
            # ---------------------------
            elif ext == "pdf":
                temp_docx = os.path.join(temp_dir, filename.replace(".pdf", ".docx"))
                doc = Document()
                with pdfplumber.open(temp_input) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            doc.add_paragraph(text)
                doc.save(temp_docx)
                converted_files.append(temp_docx)

        # ---------------------------
        # Return files to user
        # ---------------------------

        # If only one file, send directly
        if len(converted_files) == 1:
            return send_file(converted_files[0], as_attachment=True)

        # If multiple files, create ZIP
        zip_path = os.path.join(temp_dir, "converted_files.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for f in converted_files:
                if os.path.exists(f) and os.path.getsize(f) > 0:
                    zipf.write(f, arcname=os.path.basename(f))

        return send_file(zip_path, as_attachment=True)

    finally:
        # ---------------------------
        # Cleanup temporary files
        # ---------------------------
        shutil.rmtree(temp_dir, ignore_errors=True)

# -------------------------------
# Run Flask app
# -------------------------------
if __name__ == "__main__":
    app.run(debug=True)
